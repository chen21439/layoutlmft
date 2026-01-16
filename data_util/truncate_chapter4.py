"""
删除 docx 文件中 "第四章 投标文件组成要求及格式" 标题及其后面的内容，保存到指定目录

用法:
  python truncate_chapter4.py [batch_name] [file1.docx file2.docx ...]

示例:
  python truncate_chapter4.py batch1                    # 处理默认文件列表，保存到 batch1
  python truncate_chapter4.py batch2 file1.docx        # 处理指定文件，保存到 batch2
"""

import re
import sys
from pathlib import Path
from docx import Document


# 默认的文件列表
DEFAULT_FILES = [
    "修改1013深圳市宝安区燕罗街道中心幼儿园等七所公办幼儿园物业管理服务.docx",
    "批注_BACG2025000XXX_深圳市宝安中学（集团）第二实验学校物业管理服务项目A.docx",
    "批注_BAZXCG-2025-00XXX深圳市宝安中学外国语学校(集团)宝城小学教学器材采购项目.docx",
    "批注_大鹏新区建筑工务署办公楼物业管理服务项目.docx",
    "批注_[PSCG2025000096-A]深圳实验坪山学校办公及食堂家具采购项目.docx",
    "批注_深圳市坪山区中心医院物业管理服务项目招标文件.docx",
    "批注_5.19深圳市水库小学新改扩建二装图书馆设备采购招标文件.docx",
    "批注_BACG2025XXX_深圳市宝安区信息中心宝安区大数据资源中心数据运营服务（2025年）项目（A包）-1011.docx",
    "批注_A街道办事处办公物业服务管理采购.docx",
    "深圳市盐田区教育局2024年盐田区幼儿园建设-教室及办公家具设备购置.docx",
    "深圳市盐田区教育局盐田区属三校教学办公家具设备购置项目.docx",
    "批注_北京大学深圳医院内镜清洗消毒工作站和内镜洗消追溯管理系统AI采购.docx",
]

# 源文件夹路径
SOURCE_DIR = Path("/mnt/e/下载/a/客户文件")

# 输出基础目录
OUTPUT_BASE_DIR = Path("/mnt/e/models/data/Section/tender_document/docx/concat")

# 要查找的章节标题模式（正文标题）
CHAPTER_PATTERN = re.compile(r"^第四章\s*(投标文件组成要求及格式|招标文件组成要求及格式|投标文件组成要求及格式)")


def find_body_title_index(doc: Document) -> int:
    """
    找到正文标题"第四章 投标文件组成要求及格式"的段落索引

    正文标题的特征：
    - 以"第四章"开头
    - Heading样式 或 位置在10%-60%之间
    """
    total_paras = len(doc.paragraphs)

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 检查是否以"第四章"开头
        if not text.startswith("第四章"):
            continue

        # 检查位置百分比
        percent = (para_idx / total_paras) * 100 if total_paras > 0 else 0
        style = para.style.name if para.style else ""

        # 正文标题：Heading样式 或 位置在10%-60%之间
        if "Heading" in style or (10 <= percent <= 60):
            # 确认是投标/招标文件组成要求及格式
            if "投标文件" in text or "招标文件" in text:
                if "组成要求" in text or "格式" in text:
                    return para_idx

    return -1


def delete_from_paragraph(doc: Document, start_idx: int):
    """
    删除从 start_idx 开始的所有段落
    """
    # 获取文档body元素
    body = doc.element.body

    # 获取所有要删除的段落元素
    paragraphs_to_delete = doc.paragraphs[start_idx:]

    for para in paragraphs_to_delete:
        p = para._element
        p.getparent().remove(p)


def process_file(src_path: Path, dst_path: Path) -> dict:
    """
    处理单个文件：找到第四章正文标题，删除该标题及后面内容，保存
    """
    result = {
        "file": src_path.name,
        "success": False,
        "original_paragraphs": 0,
        "remaining_paragraphs": 0,
        "cut_at_index": -1,
    }

    try:
        doc = Document(src_path)
        result["original_paragraphs"] = len(doc.paragraphs)

        # 找到正文标题位置
        cut_idx = find_body_title_index(doc)

        if cut_idx == -1:
            result["error"] = "未找到正文标题"
            return result

        result["cut_at_index"] = cut_idx

        # 删除从该位置开始的所有段落
        delete_from_paragraph(doc, cut_idx)

        result["remaining_paragraphs"] = len(doc.paragraphs)

        # 保存到目标路径
        doc.save(dst_path)
        result["success"] = True
        result["output_path"] = str(dst_path)

    except Exception as e:
        result["error"] = str(e)

    return result


def main():
    # 解析命令行参数
    if len(sys.argv) > 1:
        batch_name = sys.argv[1]
    else:
        batch_name = "batch1"

    # 解析文件列表
    if len(sys.argv) > 2:
        target_files = sys.argv[2:]
    else:
        target_files = DEFAULT_FILES

    # 输出目录
    output_dir = OUTPUT_BASE_DIR / batch_name
    output_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 100)
    print("删除第四章及后续内容，保存到指定目录")
    print(f"输出目录: {output_dir}")
    print(f"文件数量: {len(target_files)}")
    print("=" * 100)
    print()

    results = []
    success_count = 0

    for filename in target_files:
        src_path = SOURCE_DIR / filename
        dst_path = output_dir / filename

        if not src_path.exists():
            print(f"[跳过] 文件不存在: {filename}")
            continue

        print(f"[处理] {filename}")
        result = process_file(src_path, dst_path)
        results.append(result)

        if result["success"]:
            success_count += 1
            print(f"  ✓ 成功: 原 {result['original_paragraphs']} 段 -> 保留 {result['remaining_paragraphs']} 段")
            print(f"    截断位置: 段落 {result['cut_at_index']}")
        else:
            print(f"  ✗ 失败: {result.get('error', '未知错误')}")
        print()

    # 汇总
    print("=" * 100)
    print(f"处理完成: 成功 {success_count}/{len(results)}")
    print(f"输出目录: {output_dir}")
    print("=" * 100)

    return results


if __name__ == "__main__":
    main()
