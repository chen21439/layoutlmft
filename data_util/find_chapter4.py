"""
在 docx 文件中查找 "第四章 招标文件组成要求及格式" 的位置信息
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


# E盘中存在但列表中没有的14个文件
TARGET_FILES = [
    "修改1013深圳市宝安区燕罗街道中心幼儿园等七所公办幼儿园物业管理服务.docx",
    "批注_BACG2025000XXX_深圳市宝安中学（集团）第二实验学校物业管理服务项目A.docx",
    "批注_BAZXCG-2025-00XXX深圳市宝安中学外国语学校(集团)宝城小学教学器材采购项目.docx",
    "批注_大鹏新区建筑工务署办公楼物业管理服务项目.docx",
    "批注_[PSCG2025000096-A]深圳实验坪山学校办公及食堂家具采购项目.docx",
    "批注_深圳市坪山区中心医院物业管理服务项目招标文件.docx",
    "批注_5.19深圳市水库小学新改扩建二装图书馆设备采购招标文件.docx",
    "批注_BACG2025XXX_深圳市宝安区信息中心宝安区大数据资源中心数据运营服务（2025年）项目（A包）-1011.docx",
    "批注_A街道办事处办公物业服务管理采购.docx",
    "深圳市盐田区教育局2024年盐田区幼儿园建设-教室及办公家具设备购置.docx",
    "深圳市盐田区教育局盐田区属三校教学办公家具设备购置项目.docx",
    "批注_北京大学深圳医院内镜清洗消毒工作站和内镜洗消追溯管理系统AI采购.docx",
]

# 源文件夹路径
SOURCE_DIR = Path("/mnt/e/下载/a/客户文件")

# 要查找的章节标题模式
CHAPTER_PATTERN = re.compile(r"第四章\s*招标文件组成要求及格式|第四章.*投标文件.*格式")


def find_chapter4_in_docx(docx_path: Path) -> dict:
    """
    在 docx 文件中查找第四章的位置信息

    返回:
        dict: 包含文件名、是否找到、段落索引、段落文本等信息
    """
    result = {
        "file": docx_path.name,
        "found": False,
        "total_paragraphs": 0,
        "locations": [],
        "middle_location": None,  # 最接近中间的位置
    }

    try:
        doc = Document(docx_path)
        total_paras = len(doc.paragraphs)
        result["total_paragraphs"] = total_paras

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 检查是否匹配第四章标题
            if CHAPTER_PATTERN.search(text) or "第四章" in text:
                percent = (para_idx / total_paras) * 100 if total_paras > 0 else 0
                location_info = {
                    "paragraph_index": para_idx,
                    "text": text[:100],  # 截取前100字符
                    "style": para.style.name if para.style else None,
                    "position_percent": round(percent, 1),
                }
                result["locations"].append(location_info)
                result["found"] = True

        # 找出最接近中间(50%)的位置
        if result["locations"]:
            middle_loc = min(result["locations"],
                           key=lambda x: abs(x["position_percent"] - 50))
            result["middle_location"] = middle_loc

    except Exception as e:
        result["error"] = str(e)

    return result


def main():
    print("=" * 100)
    print("查找 docx 文件中 '第四章 招标文件组成要求及格式' 的位置")
    print("=" * 100)
    print()

    results = []

    for filename in TARGET_FILES:
        filepath = SOURCE_DIR / filename

        if not filepath.exists():
            print(f"[跳过] 文件不存在: {filename}")
            continue

        print(f"{'=' * 100}")
        print(f"文件: {filename}")
        result = find_chapter4_in_docx(filepath)
        results.append(result)

        if result.get("error"):
            print(f"  错误: {result['error']}")
        elif result["found"]:
            print(f"总段落数: {result['total_paragraphs']}")
            print(f"找到 {len(result['locations'])} 处匹配:")
            print("-" * 100)
            print(f"{'序号':<4} {'段落':<8} {'位置%':<8} {'样式':<20} {'是否中间':<10} {'文本'}")
            print("-" * 100)

            for i, loc in enumerate(result["locations"], 1):
                is_middle = "★ 中间" if loc == result["middle_location"] else ""
                text_preview = loc["text"][:40].replace("\n", " ")
                print(f"{i:<4} {loc['paragraph_index']:<8} {loc['position_percent']:<8}% {loc['style'] or 'None':<20} {is_middle:<10} {text_preview}")

            print("-" * 100)
            mid = result["middle_location"]
            if mid:
                print(f"\n>>> 推荐的中间位置: 段落 {mid['paragraph_index']} (位置 {mid['position_percent']}%)")
                print(f"    文本: {mid['text'][:80]}")
                print(f"    样式: {mid['style']}")
        else:
            print(f"  未找到匹配")
        print()

    # 汇总统计 - 正文标题位置（Heading样式 或 位置在10-60%之间）
    print("=" * 100)
    print("汇总: 正文标题位置 (Heading样式 或 位置在10%-60%)")
    print("=" * 100)
    print(f"{'文件名':<60} {'段落':<8} {'位置%':<10} {'样式':<15} {'状态'}")
    print("-" * 100)

    for r in results:
        fname = r["file"][:55] + "..." if len(r["file"]) > 58 else r["file"]

        if r.get("error"):
            print(f"{fname:<60} {'错误':<8}")
            continue

        if not r.get("locations"):
            print(f"{fname:<60} {'-':<8} {'-':<10} {'-':<15} 未找到第四章")
            continue

        # 筛选正文标题：Heading样式 或 位置在10-60%之间
        body_titles = [
            loc for loc in r["locations"]
            if (loc["style"] and "Heading" in loc["style"])
            or (10 <= loc["position_percent"] <= 60)
        ]

        # 进一步筛选：只保留真正的章节标题（以"第四章"开头）
        body_titles = [
            loc for loc in body_titles
            if loc["text"].strip().startswith("第四章")
        ]

        if body_titles:
            # 取第一个（最靠前的正文标题）
            title = body_titles[0]
            print(f"{fname:<60} {title['paragraph_index']:<8} {title['position_percent']:<10}% {title['style'] or 'Normal':<15} ✓ 有")
        else:
            print(f"{fname:<60} {'-':<8} {'-':<10} {'-':<15} ✗ 无正文标题")

    return results


if __name__ == "__main__":
    main()
