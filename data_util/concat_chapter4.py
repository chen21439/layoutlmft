"""
拼接裁切后的docx和文档A的"第四章..."后的内容

功能:
  --docx     拼接docx文件
  --json     拼接json文件（包含重新挂载非section节点）
  --remount  仅重新挂载非section节点（不拼接文档A）

用法:
  python concat_chapter4.py batch1 --docx          # 只拼接docx
  python concat_chapter4.py batch1 --json          # 只拼接json
  python concat_chapter4.py batch1 --docx --json   # 两者都做
  python concat_chapter4.py batch1 --remount       # 仅重新挂载，不拼接
  python concat_chapter4.py --dir /path/to/dir --json  # 直接指定目录

目录结构:
  concat/batch1/
  ├── *.docx           ← 裁切后的docx
  ├── *.json           ← 你修正后的json (section parent_id已修复)
  └── fulltext/        ← 输出目录
      ├── *.docx       ← 拼接后的完整docx
      └── *.json       ← 最终json
"""

import argparse
import json
import re
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_BREAK


# 文档A的路径 (用于拼接"第四章..."后的内容)
DOC_A_DOCX = Path("/mnt/e/models/data/Section/tender_document/docx/concat/src/批注_[GMCG2025000068-A]深圳市光明区机关事务管理中心光明区档案综合服务中心物业管理服务.docx")
DOC_A_JSON = Path("/mnt/e/models/data/Section/tender_document/docx/concat/src/批注_[GMCG2025000068-A]深圳市光明区机关事务管理中心光明区档案综合服务中心物业管理服务.json")


# ============================================================================
# 查找第四章位置
# ============================================================================

def find_chapter4_paragraph_index(doc: Document) -> int:
    """在docx中找到正文标题"第四章 投标文件组成要求及格式"的段落索引"""
    total_paras = len(doc.paragraphs)

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or not text.startswith("第四章"):
            continue

        percent = (para_idx / total_paras) * 100 if total_paras > 0 else 0
        style = para.style.name if para.style else ""

        if "Heading" in style or (10 <= percent <= 60):
            if "投标文件" in text or "招标文件" in text:
                if "组成要求" in text or "格式" in text:
                    return para_idx
    return -1


def find_chapter4_json_index(line_json: list) -> int:
    """在line_json中找到"第四章 投标文件组成要求及格式"的索引"""
    for idx, item in enumerate(line_json):
        text = item.get("text", "").strip()
        cls = item.get("class", "")

        if cls == "section" and text.startswith("第四章"):
            if "投标文件" in text or "招标文件" in text:
                if "组成要求" in text or "格式" in text:
                    return idx
    return -1


# ============================================================================
# 重新挂载非section节点
# ============================================================================

def extract_data_list(json_data):
    """从JSON中提取dataList，支持两种格式"""
    if isinstance(json_data, list):
        return json_data
    elif isinstance(json_data, dict):
        # 格式: {"data": {"dataList": [...]}}
        if "data" in json_data and "dataList" in json_data["data"]:
            return json_data["data"]["dataList"]
        # 格式: {"dataList": [...]}
        elif "dataList" in json_data:
            return json_data["dataList"]
    return None


def wrap_data_list(data_list, original_json):
    """将dataList包装回原始格式"""
    if isinstance(original_json, list):
        return data_list
    elif isinstance(original_json, dict):
        result = deepcopy(original_json)
        if "data" in result and "dataList" in result["data"]:
            result["data"]["dataList"] = data_list
        elif "dataList" in result:
            result["dataList"] = data_list
        return result
    return data_list


def remount_non_section_nodes(line_json: list) -> list:
    """
    根据修正后的section，重新挂载非section节点

    规则:
    - section: 保持原有的 parent_id 和 relation
    - fstline:
      - 第一个: contain 到所属 section
      - 后续: equality 到前一个 fstline
    - paraline: connect 到前一个节点
    - table: contain 到所属 section
    - 其他: contain 到所属 section
    """
    result = deepcopy(line_json)

    # 按 line_id 排序（阅读顺序）
    result.sort(key=lambda x: x.get("line_id", x.get("id", 0)))

    current_section_line_id = None
    last_fstline_line_id = None
    last_node_line_id = None

    for item in result:
        cls = item.get("class", "")
        line_id = item.get("line_id", item.get("id"))

        if cls == "section":
            current_section_line_id = line_id
            last_fstline_line_id = None
            last_node_line_id = line_id

        elif cls == "fstline":
            if last_fstline_line_id is None:
                if current_section_line_id is not None:
                    item["parent_id"] = current_section_line_id
                    item["relation"] = "contain"
            else:
                item["parent_id"] = last_fstline_line_id
                item["relation"] = "equality"

            last_fstline_line_id = line_id
            last_node_line_id = line_id

        elif cls == "paraline":
            if last_node_line_id is not None:
                item["parent_id"] = last_node_line_id
                item["relation"] = "connect"
            last_node_line_id = line_id

        elif cls == "table":
            if current_section_line_id is not None:
                item["parent_id"] = current_section_line_id
                item["relation"] = "contain"
            last_node_line_id = line_id

        else:
            if current_section_line_id is not None:
                item["parent_id"] = current_section_line_id
                item["relation"] = "contain"
            last_node_line_id = line_id

    return result


# ============================================================================
# DOCX 拼接
# ============================================================================

def add_page_break(doc: Document):
    """在文档末尾添加分页符"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def copy_paragraph(source_para, target_doc):
    """复制段落到目标文档"""
    new_para = target_doc.add_paragraph()

    if source_para.style:
        try:
            new_para.style = source_para.style.name
        except:
            pass

    if source_para.paragraph_format:
        pf = new_para.paragraph_format
        spf = source_para.paragraph_format
        if spf.alignment:
            pf.alignment = spf.alignment
        if spf.left_indent:
            pf.left_indent = spf.left_indent
        if spf.right_indent:
            pf.right_indent = spf.right_indent
        if spf.first_line_indent:
            pf.first_line_indent = spf.first_line_indent
        if spf.space_before:
            pf.space_before = spf.space_before
        if spf.space_after:
            pf.space_after = spf.space_after
        if spf.line_spacing:
            pf.line_spacing = spf.line_spacing

    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        if run.bold:
            new_run.bold = run.bold
        if run.italic:
            new_run.italic = run.italic
        if run.underline:
            new_run.underline = run.underline
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name

    return new_para


def copy_table(source_table, target_doc):
    """复制表格到目标文档"""
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    new_table = target_doc.add_table(rows=rows, cols=cols)

    if source_table.style:
        try:
            new_table.style = source_table.style
        except:
            pass

    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_table.rows[i].cells[j].text = cell.text

    return new_table


def concat_docx(truncated_docx_path: Path, doc_a_docx_path: Path, output_path: Path) -> bool:
    """拼接docx文件"""
    try:
        truncated_doc = Document(truncated_docx_path)
        doc_a = Document(doc_a_docx_path)

        chapter4_idx = find_chapter4_paragraph_index(doc_a)
        if chapter4_idx == -1:
            print(f"  [错误] 文档A中未找到第四章")
            return False

        add_page_break(truncated_doc)

        body_elements = list(doc_a.element.body)
        para_count = 0
        copying = False

        for element in body_elements:
            if element.tag.endswith('p'):
                if para_count == chapter4_idx:
                    copying = True

                if copying:
                    for para in doc_a.paragraphs:
                        if para._element == element:
                            copy_paragraph(para, truncated_doc)
                            break
                para_count += 1

            elif element.tag.endswith('tbl') and copying:
                for table in doc_a.tables:
                    if table._element == element:
                        copy_table(table, truncated_doc)
                        break

        output_path.parent.mkdir(parents=True, exist_ok=True)
        truncated_doc.save(output_path)
        return True

    except Exception as e:
        print(f"  [错误] docx拼接失败: {e}")
        return False


# ============================================================================
# JSON 拼接
# ============================================================================

def concat_json(truncated_json: list, doc_a_json_path: Path) -> list:
    """
    拼接line_json
    1. truncated_json 已经过 remount_non_section_nodes 修正
    2. 从文档A的json中提取"第四章..."及后续元素
    3. 重新编排id, line_id, parent_id, page
    """
    with open(doc_a_json_path, 'r', encoding='utf-8') as f:
        doc_a_raw = json.load(f)

    # 提取dataList
    doc_a_json = extract_data_list(doc_a_raw)
    if doc_a_json is None:
        doc_a_json = doc_a_raw  # 尝试直接使用

    chapter4_idx = find_chapter4_json_index(doc_a_json)
    if chapter4_idx == -1:
        print(f"  [错误] 文档A的JSON中未找到第四章")
        return None

    def get_page(item):
        """从 item 中获取 page，支持两种格式"""
        # 格式1: location: [{"page": 24, ...}]
        loc = item.get("location", [])
        if loc and isinstance(loc, list) and len(loc) > 0:
            return int(loc[0].get("page", 0))
        # 格式2: "page": "48" 或 "page": 48
        page = item.get("page", 0)
        return int(page) if page else 0

    if truncated_json:
        last_id = max(item.get("id", 0) for item in truncated_json)
        last_line_id = max(item.get("line_id", 0) for item in truncated_json)
        last_page = max(get_page(item) for item in truncated_json)
    else:
        last_id = -1
        last_line_id = -1
        last_page = -1

    chapter4_elements = doc_a_json[chapter4_idx:]
    chapter4_start_page = int(chapter4_elements[0].get("page", 0))

    page_offset = last_page + 1 - chapter4_start_page

    # line_id 映射: old_line_id -> new_line_id
    line_id_mapping = {}
    chapter4_first_line_id = chapter4_elements[0].get("line_id", 0)

    # 第一遍：建立 line_id 映射
    for i, item in enumerate(chapter4_elements):
        old_line_id = item.get("line_id", item.get("id", 0))
        new_line_id = last_line_id + 1 + i
        line_id_mapping[old_line_id] = new_line_id

    # 第二遍：生成新元素
    new_elements = []
    for i, item in enumerate(chapter4_elements):
        new_item = {}

        old_line_id = item.get("line_id", item.get("id", 0))
        old_parent_id = item.get("parent_id", -1)
        old_page = int(item.get("page", 0))
        old_box = item.get("box", [0, 0, 0, 0])

        new_line_id = line_id_mapping[old_line_id]
        new_page = old_page + page_offset

        # 基本字段
        new_item["line_id"] = new_line_id
        new_item["id"] = new_line_id

        # parent_id 映射
        if i == 0:
            # 第四章本身的 parent_id 设为 -1
            new_item["parent_id"] = -1
        elif old_parent_id < chapter4_first_line_id:
            # 指向第四章之前的元素，设为 -1
            new_item["parent_id"] = -1
        elif old_parent_id in line_id_mapping:
            # 映射到新的 line_id
            new_item["parent_id"] = line_id_mapping[old_parent_id]
        else:
            new_item["parent_id"] = -1

        # relation
        new_item["relation"] = item.get("relation", "contain")

        # text
        new_item["text"] = item.get("text", "")

        # class
        new_item["class"] = item.get("class", "")

        # location: 将 page + box 转换为 location 格式
        new_item["location"] = [{
            "page": new_page,
            "l": float(old_box[0]) if len(old_box) > 0 else 0.0,
            "t": float(old_box[1]) if len(old_box) > 1 else 0.0,
            "r": float(old_box[2]) if len(old_box) > 2 else 0.0,
            "b": float(old_box[3]) if len(old_box) > 3 else 0.0,
            "coord_origin": "TOPLEFT"
        }]

        # is_section (如果有)
        if "is_section" in item:
            new_item["is_section"] = item.get("is_section")
        elif item.get("class") == "section":
            new_item["is_section"] = True
        else:
            new_item["is_section"] = False

        # is_meta (如果有)
        if "is_meta" in item:
            new_item["is_meta"] = item.get("is_meta")

        new_elements.append(new_item)

    result_json = truncated_json + new_elements
    return result_json


# ============================================================================
# 路径转换
# ============================================================================

def convert_windows_path_to_wsl(path_str: str) -> str:
    """
    将 Windows 路径转换为 WSL 路径
    E:\\models\\data -> /mnt/e/models/data
    E:modelsdata (bash吃掉反斜杠后) -> /mnt/e/models/data
    """
    if not path_str:
        return path_str

    # 检查是否是 Windows 路径格式 (盘符:)
    if len(path_str) >= 2 and path_str[1] == ':':
        drive_letter = path_str[0].lower()
        # 移除盘符和冒号
        rest_path = path_str[2:]
        # 转换反斜杠为正斜杠
        rest_path = rest_path.replace('\\', '/')
        # 移除开头的斜杠（如果有）
        rest_path = rest_path.lstrip('/')
        # 构建 WSL 路径
        return f"/mnt/{drive_letter}/{rest_path}"

    return path_str


# ============================================================================
# 主函数
# ============================================================================

def process_docx(batch_dir: Path, output_dir: Path):
    """处理所有docx文件"""
    print("\n[DOCX拼接]")
    docx_files = list(batch_dir.glob("*.docx"))
    success_count = 0

    for docx_path in docx_files:
        output_path = output_dir / docx_path.name
        print(f"  {docx_path.name} ... ", end="")

        if concat_docx(docx_path, DOC_A_DOCX, output_path):
            print("✓")
            success_count += 1
        else:
            print("✗")

    print(f"  完成: {success_count}/{len(docx_files)}")
    return success_count


def process_json_remount(batch_dir: Path, output_dir: Path):
    """仅重新挂载非section节点"""
    print("\n[JSON重新挂载] (仅修正非section节点)")

    json_files = list(batch_dir.glob("*.json"))
    success_count = 0

    for json_path in json_files:
        output_path = output_dir / json_path.name
        print(f"  {json_path.name} ... ", end="")

        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                original_json = json.load(f)

            data_list = extract_data_list(original_json)
            if data_list is None:
                print("✗ (无法解析JSON格式)")
                continue

            original_count = len(data_list)

            # 重新挂载非section节点
            final_list = remount_non_section_nodes(data_list)

            output_dir.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(final_list, f, ensure_ascii=False, indent=2)
            print(f"✓ ({original_count})")
            success_count += 1

        except Exception as e:
            print(f"✗ ({e})")

    print(f"  完成: {success_count}/{len(json_files)}")
    return success_count


def process_json_concat(batch_dir: Path, output_dir: Path):
    """只拼接文档A的第四章后内容（不修改前面的）"""
    print("\n[JSON拼接] (只拼接第四章后内容)")

    json_files = list(batch_dir.glob("*.json"))
    success_count = 0

    for json_path in json_files:
        output_path = output_dir / json_path.name
        print(f"  {json_path.name} ... ", end="")

        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                original_json = json.load(f)

            data_list = extract_data_list(original_json)
            if data_list is None:
                print("✗ (无法解析JSON格式)")
                continue

            original_count = len(data_list)

            # 直接拼接，不修改前面的内容
            final_list = concat_json(data_list, DOC_A_JSON)
            if final_list is None:
                print("✗ (拼接失败)")
                continue

            output_dir.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(final_list, f, ensure_ascii=False, indent=2)
            print(f"✓ ({original_count} → {len(final_list)})")
            success_count += 1

        except Exception as e:
            print(f"✗ ({e})")

    print(f"  完成: {success_count}/{len(json_files)}")
    return success_count


def main():
    parser = argparse.ArgumentParser(description="拼接裁切后的docx和文档A的第四章后内容")
    parser.add_argument("batch", nargs="?", help="批次名称，如 batch1 (当不使用 --dir 时必需)")
    parser.add_argument("--dir", type=str, help="直接指定工作目录路径（如果指定，将忽略 batch 参数）")
    parser.add_argument("--docx", action="store_true", help="拼接docx文件")
    parser.add_argument("--json", action="store_true", help="只拼接第四章后的json（不修改前面的）")
    parser.add_argument("--remount", action="store_true", help="仅重新挂载非section节点（不拼接）")

    args = parser.parse_args()

    # 如果没有指定任何操作，显示帮助
    if not args.docx and not args.json and not args.remount:
        parser.print_help()
        print("\n请至少指定一个操作: --docx, --json, 或 --remount")
        return

    # 目录配置
    if args.dir:
        # 使用 --dir 直接指定目录，自动转换 Windows 路径
        converted_path = convert_windows_path_to_wsl(args.dir)
        batch_dir = Path(converted_path)
    elif args.batch:
        # 检查 batch 参数是否看起来像 Windows 路径（包含盘符:）
        if len(args.batch) >= 2 and args.batch[1] == ':':
            # 看起来是 Windows 路径，直接转换
            converted_path = convert_windows_path_to_wsl(args.batch)
            batch_dir = Path(converted_path)
            print(f"\n[提示] 检测到 Windows 路径，已自动转换")
            print(f"  原始: {args.batch}")
            print(f"  转换: {converted_path}")
            print(f"  建议: 在 bash 中使用 Windows 路径时请加引号，如: --dir \"E:\\path\\to\\dir\"\n")
        else:
            # 使用 batch 参数构建目录
            batch_dir = Path(f"/mnt/e/models/data/Section/tender_document/docx/concat/{args.batch}")
    else:
        parser.print_help()
        print("\n[错误] 请指定 batch 参数或使用 --dir 指定目录")
        return

    output_dir = batch_dir / "fulltext"

    print("=" * 80)
    # 根据操作类型显示不同标题
    if args.remount:
        print("重新挂载非section节点")
    elif args.docx and args.json:
        print("拼接docx和json文件")
    elif args.docx:
        print("拼接docx文件")
    elif args.json:
        print("拼接json文件")
    print("=" * 80)
    print(f"工作目录: {batch_dir}")
    print(f"输出目录: {output_dir}")

    if not batch_dir.exists():
        print(f"\n[错误] Batch目录不存在: {batch_dir}")
        return

    output_dir.mkdir(parents=True, exist_ok=True)

    # 执行操作
    if args.docx:
        process_docx(batch_dir, output_dir)

    if args.remount:
        process_json_remount(batch_dir, output_dir)

    if args.json:
        process_json_concat(batch_dir, output_dir)

    print("\n" + "=" * 80)
    print("完成")
    print("=" * 80)


if __name__ == "__main__":
    main()
